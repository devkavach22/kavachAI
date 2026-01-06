
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from PyPDF2 import PdfReader
from docx import Document
import pytesseract
import io
from pdf2image import convert_from_bytes
import fitz  # PyMuPDF
from PIL import Image
from services.CV_parsing_chain import get_cv_data_from_openrouter_model
router = APIRouter()


def extract_images_from_pdf(file_bytes):
    images = []
    pdf_file = fitz.open(stream=file_bytes, filetype="pdf")
    for page_index in range(len(pdf_file)):
        page = pdf_file.load_page(page_index)
        image_list = page.get_images(full=True)
        # Optionally print: print(f"[+] Found a total of {len(image_list)} images on page {page_index}")
        for image_index, img in enumerate(image_list, start=1):
            xref = img[0]
            base_image = pdf_file.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image = Image.open(io.BytesIO(image_bytes))
            images.append({
                "page": page_index + 1,
                "image_index": image_index,
                "image": image,
                "extension": image_ext
            })
    return images

def extract_images_from_docx(file_bytes):
    images = []
    doc = Document(io.BytesIO(file_bytes))
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            img = Image.open(io.BytesIO(image_data))
            images.append(img)
    return images



@router.post("/uploadCV")
async def upload_cv(name: str = Form(...), cv: UploadFile = File(...)):
    # Allowed file types
    allowed_types = ["application/pdf", 
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]


    if cv.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Only PDF or Word (.docx) files are allowed.")

    # Read file content
    file_bytes = await cv.read()

    extracted_text = ""

    try:
        # --- Extract text from PDF ---
        if cv.content_type == "application/pdf":
            pdf_reader = PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() or ""

            # If no text extracted, extract all images from PDF and run OCR
            if not extracted_text.strip():
                try:
                    images = extract_images_from_pdf(file_bytes)
                    for img_info in images:
                        extracted_text += pytesseract.image_to_string(img_info["image"])
                except Exception as ocr_e:
                    # Fallback: try OCR on page images (as before)
                    try:
                        images = convert_from_bytes(file_bytes)
                        for img in images:
                            extracted_text += pytesseract.image_to_string(img)
                    except Exception as ocr_e2:
                        # Poppler not installed or other error
                        raise HTTPException(status_code=500, detail=f"Error extracting text with OCR: {str(ocr_e)}; {str(ocr_e2)}. If using pdf2image, ensure poppler is installed and in your PATH.")

        # --- Extract text from DOCX ---
        elif cv.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"

            # If no text extracted, extract all images from DOCX and run OCR
            if not extracted_text.strip():
                try:
                    images = extract_images_from_docx(file_bytes)
                    for img in images:
                        extracted_text += pytesseract.image_to_string(img)
                except Exception as ocr_e:
                    raise HTTPException(status_code=500, detail=f"Error extracting text with OCR from DOCX: {str(ocr_e)}")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
    if not extracted_text.strip():
        raise HTTPException(status_code=400, detail="Could not extract any text from the uploaded file.")

    # Process the extracted text to get structured data
    try:
        structured_data = get_cv_data_from_openrouter_model(extracted_text)
        return {
        "message": "File processed successfully!",
        "file_name": cv.filename,
        "file_type": cv.content_type,
        "name": name,
        "structured_data": structured_data
    }
    except Exception as e:
        print("Error processing CV data:", str(e))
        raise HTTPException(status_code=500, detail=f"Error processing CV data: {str(e)}")


# @router.get("/uploadTest")
# async def upload_test():
#     name = "Test User"
#     extracted_text = "This is a test extracted text."
#     return {
#         "message": "Test endpoint working!",
#         "file_name": "testfile.pdf",
#         "file_type": "application/pdf",
#         "name": name,
#         "extracted_text": extracted_text.strip()
#     }