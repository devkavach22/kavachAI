import io
import json
import os
import re
import sqlite3
import tempfile
import zipfile
from typing import Any, Dict
from pathlib import Path

import docx2txt
import easyocr
import fitz  # PyMuPDF
import pandas as pd
import rarfile
import xlrd
from bs4 import BeautifulSoup
from docx import Document
from fastapi import (
    APIRouter,
    File,
    HTTPException,
    UploadFile,
    WebSocket,
    WebSocketDisconnect,
)
from pdf2image import convert_from_bytes
from PIL import Image
from pydantic import BaseModel, Field
from typing import Literal, Optional

from services.data_extraction import FileDataExtractor
from pipeline.home_chat_data_extraction import (
    FileDataExtractor as pipeline_FileDataExtractor,
)
from services.Home_page_chain import run_home_page_chain

router = APIRouter()


def extract_images_from_pdf(file_bytes):
    images = []
    pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
    for page_index in range(len(pdf_doc)):
        page = pdf_doc.load_page(page_index)
        pix = page.get_pixmap(dpi=200)  # Render page to image
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        images.append(img)
    return images


def extract_images_from_docx(file_bytes):
    images = []
    doc = Document(io.BytesIO(file_bytes))
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            img = Image.open(io.BytesIO(image_data))
            images.append(img)
    return images


def extract_text_from_image(images):
    reader = easyocr.Reader(["en"])
    # result = reader.readtext(image_bytes, detail=0, paragraph=True)
    # return " ".join(result)
    all_OCR_text = []
    for img in images:
        with io.BytesIO() as buf:
            img.save(buf, format="PNG")
            img_bytes = buf.getvalue()
            result = reader.readtext(img_bytes, detail=0, paragraph=True)
            all_OCR_text.extend(result)
    return " ".join(all_OCR_text)


@staticmethod
async def _extract_zip(file_bytes: bytes, ext: str) -> Dict[str, Any]:
    """
    Extract and process all files inside a ZIP archive.
    Automatically detects file types and calls FileDataExtractor recursively.
    Returns only processed results (no metadata).
    """
    processed_results = {}

    with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
        files = zf.namelist()

        for filename in files:
            # Skip directories
            if filename.endswith("/"):
                continue

            with zf.open(filename) as inner_file:
                inner_file_bytes = inner_file.read()

            # Detect file extension
            inner_ext = os.path.splitext(filename)[1].lower()

            # Initialize extractor
            extractor = FileDataExtractor()

            # Supported extraction methods
            extraction_methods = {
                ".pdf": extractor._extract_pdf,
                ".jpg": extractor._extract_image,
                ".jpeg": extractor._extract_image,
                ".png": extractor._extract_image,
                ".bmp": extractor._extract_image,
                ".tiff": extractor._extract_image,
                ".docx": extractor._extract_docx,
                ".xls": extractor._extract_excel,
                ".xlsx": extractor._extract_excel,
                ".xlsm": extractor._extract_excel,
                ".xltx": extractor._extract_excel,
                ".xltm": extractor._extract_excel,
                ".json": extractor._extract_json,
                ".db": extractor._extract_sql,
                ".sqlite": extractor._extract_sql,
                ".sql": extractor._extract_sql,
                ".html": extractor._extract_html,
                ".txt": extractor._extract_txt,
                ".zip": extractor._extract_zip,  # support nested ZIPs
                ".mp3": extractor._extract_audio,
                ".wav": extractor._extract_audio,
                # 🎥 NEW VIDEO SUPPORT
                ".mp4": extractor._extract_video,
                ".avi": extractor._extract_video,
                ".mov": extractor._extract_video,
                ".mkv": extractor._extract_video,
                ".webm": extractor._extract_video,
                ".flv": extractor._extract_video,
                ".wmv": extractor._extract_video,
            }

            # Process inner file if supported
            if inner_ext in extraction_methods:
                try:
                    result = await extraction_methods[inner_ext](
                        inner_file_bytes, inner_ext
                    )
                    processed_results[filename] = result
                except Exception as e:
                    processed_results[filename] = {"error": str(e)}
            else:
                processed_results[filename] = {
                    "warning": f"Unsupported inner file type: {inner_ext}",
                }

    return {
        "type": "zip",
        "contained_files": files,
        "processed_results": processed_results,
    }


@router.post("/extract_file_data/")
async def extract_file_data(file: UploadFile = File(...)):
    """
    Universal file data extractor supporting multiple formats:
    PDF, Image, Word, Excel, JSON, SQL, ZIP, RAR, HTML, TXT, CV,MP3,WAV,MP4,AVI,MOV,WMV,FLV,WMA
    """

    filename = file.filename
    file_bytes = await file.read()
    ext = os.path.splitext(filename)[1].lower()
    extractor = FileDataExtractor()

    # --- Helper: EasyOCR Reader (for image/pdf OCR)
    reader = easyocr.Reader(["en"])

    try:
        # ========== 1️⃣ PDF ==========
        if ext == ".pdf":
            try:
                # --- Attempt 1: Direct text extraction with PyMuPDF ---
                text = ""
                with fitz.open(stream=io.BytesIO(file_bytes), filetype="pdf") as pdf:
                    for page in pdf:
                        text += page.get_text("text")

                # --- If text is minimal, fallback to OCR ---
                if len(text.strip()) < 100:  # Threshold to detect scanned PDFs
                    images = extract_images_from_pdf(file_bytes)
                    images_text = extract_text_from_image(images)
                    result_text = text + images_text
                return {"type": "pdf", "text": result_text.strip()}
            except Exception as e:
                raise HTTPException(500, f"PDF extraction failed: {e}")

        # ========== 2️⃣ IMAGE ==========
        elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
            img = Image.open(io.BytesIO(file_bytes))
            with io.BytesIO() as buf:
                img.save(buf, format="PNG")
                img_bytes = buf.getvalue()
            result = reader.readtext(img_bytes, detail=0, paragraph=True)
            return {"type": "image", "text": " ".join(result)}

        # ========== 3️⃣ WORD (.docx) ==========
        elif ext == ".docx":
            # doc = Document(io.BytesIO(file_bytes))
            # for para in doc.paragraphs:
            #     extracted_text += para.text + "\n"
            text = docx2txt.process(io.BytesIO(file_bytes))
            images = extract_images_from_docx(file_bytes)
            all_OCR_text = extract_text_from_image(images)
            text += all_OCR_text
            return {"type": "word", "text": text}

        # ========== 4️⃣ EXCEL (.xls, .xlsx) ==========
        elif ext in [".xls", ".xlsx", ".xlsm", ".xltx", ".xltm"]:
            try:
                with io.BytesIO(file_bytes) as buf:
                    if ext in [".xlsx", ".xlsm", ".xltx", ".xltm"]:
                        # ✅ Use openpyxl for modern Excel files
                        xls = pd.read_excel(buf, sheet_name=None, engine="openpyxl")

                    elif ext == ".xls":
                        # ✅ Use xlrd for legacy Excel files
                        try:
                            xls = pd.read_excel(buf, sheet_name=None, engine="xlrd")
                        except ImportError:
                            raise HTTPException(
                                500,
                                "Reading .xls files requires 'xlrd'. Install with 'pip install xlrd>=2.0.1'.",
                            )
                    else:
                        raise HTTPException(400, "Unsupported Excel format")
                result = {
                    sheet: data.to_dict(orient="records") for sheet, data in xls.items()
                }
                return {"type": "excel", "sheets": result}
            except Exception as e:
                raise HTTPException(500, f"Excel extraction failed: {e}")

        # ========== 5️⃣ JSON ==========
        elif ext == ".json":
            try:
                json_data = json.loads(file_bytes.decode("utf-8"))
                return {"type": "json", "data": json_data}
            except Exception as e:
                raise HTTPException(400, f"Invalid JSON: {e}")

        # ========== 6️⃣ SQL (SQLite .db or .sql) ==========
        elif ext in [".db", ".sqlite", ".sql"]:
            try:
                if ext in [".sql"]:
                    SQL_query = (
                        file_bytes.decode("utf-8")
                        .encode("utf-8")
                        .decode("unicode_escape")
                    )
                    # SQL_query = SQL_query.replace("\n", " ").strip()
                    return {"type": "sql_query", "query": SQL_query}

            except Exception as e:
                raise HTTPException(400, f"SQL extraction failed: {e}")

        # ========== 7️⃣ ZIP / RAR ==========
        elif ext == ".zip":
            final_result = await _extract_zip(
                file_bytes, os.path.splitext(filename)[1].lower()
            )

            return {
                "type": "zip",
                # "contained_files": files,
                "file_details": final_result,
            }

        elif ext == ".rar":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".rar") as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            with rarfile.RarFile(tmp_path) as rf:
                files = rf.namelist()
            os.remove(tmp_path)
            return {"type": "rar", "contained_files": files}

        # ========== 8️⃣ HTML ==========
        elif ext == ".html":
            soup = BeautifulSoup(
                file_bytes.decode("utf-8", errors="ignore"), "html.parser"
            )
            text = soup.get_text(separator="\n")
            links = [a["href"] for a in soup.find_all("a", href=True)]
            images = [img["src"] for img in soup.find_all("img", src=True)]
            return {
                "type": "html",
                "text": text.strip(),
                "links": links,
                "images": images,
            }

        # ========== 9️⃣ TXT ==========
        elif ext == ".txt":
            text = file_bytes.decode("utf-8", errors="ignore")
            return {"type": "txt", "text": text}

        # ========== 10️⃣ MP3 / WAV ==========
        elif ext in [".mp3", ".wav"]:
            text = await extractor._extract_audio(file_bytes, ext)
            return {"type": "audio", "text": text}

        # ========== 11️⃣ MP4 / AVI / MOV / MKV / WEBM / FLV / WMV / WMA ==========
        elif ext in [".mp4", ".avi", ".mov", ".mkv", ".webm", ".flv", ".wmv", ".wma"]:
            text = await extractor._extract_video(file_bytes, ext)
            return {"type": "video", "text": text}

        # ========== 12️⃣ UNKNOWN FILE ==========
        else:
            raise HTTPException(400, f"Unsupported file type: {ext}")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")


import logging

from pydantic import ValidationError

import json


from pydantic import BaseModel,field_validator
from typing import Optional


class FileMeta(BaseModel):
    filename: Optional[str] = None
    content_type: Optional[str] = None
    size: Optional[int] = None


class QueryMessage(BaseModel):
    user_query: Optional[str] = None
    file_meta: Optional[FileMeta] = None
    file_bytes: Optional[bytes] = None

    @field_validator("file_bytes", mode="before")
    def validate_file_bytes_meta(cls, value):
        if (value is not None) and (cls.file_meta is None):
            raise ValueError("file meta is required when file bytes is sending or provided")
        return value


logger = logging.getLogger("websocket")


@router.websocket("/ws/chat-file")
async def websocket_chat_file(websocket: WebSocket):
    await websocket.accept()
    logger.info("🟢 WebSocket connection accepted")

    extractor = pipeline_FileDataExtractor()

    try:
        while True:
            message = await websocket.receive()

            user_query: Optional[str] = None
            file_meta: Optional[FileMeta] = None
            file_bytes = bytearray()
            extracted_text: Optional[str] = None

            # 🔹 JSON message (query + file_meta together)
            if message.get("text"):
                try:
                    data = QueryMessage.model_validate_json(message["text"])
                    user_query = data.user_query
                    file_meta = data.file_meta
                except ValidationError as ve:
                    await websocket.send_json(
                        {"error": f"Invalid message format: {ve}"}
                    )
                    continue

                if user_query:
                    logger.info(f"📝 User query received: {user_query}")

                if file_meta:
                    logger.info(f"📄 File meta received: {file_meta}")
                    # 🔹 If file_meta is present, wait for file bytes
                    bytes_message = await websocket.receive()
                    if bytes_message.get("bytes"):
                        file_bytes.extend(bytes_message["bytes"])
                        logger.info(
                            f"📦 Received {len(file_bytes)} bytes for {file_meta.filename}"
                        )

                        # Extract text from file
                        ext = (
                            os.path.splitext(file_meta.filename)[1].lower()
                            if file_meta.filename
                            else ""
                        )
                        extracted_text = await extractor.extract_data(
                            bytes(file_bytes), ext
                        )
                    else:
                        logger.warning(
                            "⚠️ Expected file bytes but received something else"
                        )

            # 🔹 Process with LangChain if we have query or data
            if user_query or extracted_text:
                await websocket.send_json({"status": "processing"})

                try:
                    # Run the chain (non-streaming based on user request to use run_home_page_chain)
                    response = await run_home_page_chain(
                        user_query=user_query, file_data=extracted_text
                    )

                    await websocket.send_json({"type": "content", "content": response})

                    await websocket.send_json({"type": "done", "status": "success"})

                except Exception as e:
                    logger.error(f"🔴 Chain error: {e}")
                    await websocket.send_json(
                        {"error": f"AI processing failed: {str(e)}"}
                    )

            else:
                await websocket.send_json({"error": "No query or file data provided"})

    except WebSocketDisconnect:
        logger.info("🔴 WebSocket disconnected")
    except Exception as e:
        logger.error(f"🔴 WebSocket error: {e}")
        try:
            await websocket.close()
        except:
            pass
